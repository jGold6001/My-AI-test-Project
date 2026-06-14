import os
import time
from pathlib import Path
from typing import List

import docx
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from langchain_core.documents import Document
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pydantic import BaseModel
from pypdf import PdfReader
from qdrant_client import QdrantClient, models

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "myownai_docs")
LLM_MODEL = os.getenv("LLM_MODEL", "llama3")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")
VECTOR_SIZE = int(os.getenv("VECTOR_SIZE", "768"))
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="MyOwnAI Test Project")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class AskRequest(BaseModel):
    question: str


class IngestRequest(BaseModel):
    texts: List[str]


embeddings = OllamaEmbeddings(
    model=EMBEDDING_MODEL,
    base_url=OLLAMA_BASE_URL,
)

client = QdrantClient(url=QDRANT_URL)


def wait_for_qdrant(retries: int = 20, delay_seconds: float = 1.0) -> None:
    """Wait until Qdrant HTTP API is ready."""
    last_error = None

    for _ in range(retries):
        try:
            client.get_collections()
            return
        except Exception as exc:
            last_error = exc
            time.sleep(delay_seconds)

    raise RuntimeError(f"Qdrant is not ready: {last_error}")


def ensure_collection_exists() -> None:
    """Create the Qdrant collection on first startup if it does not exist yet."""
    wait_for_qdrant()

    if client.collection_exists(collection_name=COLLECTION_NAME):
        return

    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=models.VectorParams(
            size=VECTOR_SIZE,
            distance=models.Distance.COSINE,
        ),
    )


ensure_collection_exists()

vector_store = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name=COLLECTION_NAME,
    url=QDRANT_URL,
)

llm = ChatOllama(
    model=LLM_MODEL,
    base_url=OLLAMA_BASE_URL,
    temperature=0.2,
)


def read_uploaded_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return file_path.read_text(encoding="utf-8")

    if suffix == ".pdf":
        reader = PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if suffix == ".docx":
        document = docx.Document(str(file_path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Use .txt, .pdf, or .docx",
    )


def split_text_to_documents(text: str, source: str) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
    )

    chunks = splitter.split_text(text)

    return [
        Document(
            page_content=chunk,
            metadata={"source": source},
        )
        for chunk in chunks
        if chunk.strip()
    ]


@app.get("/health")
def health():
    return {
        "status": "ok",
        "stack": ["React", "FastAPI", "Ollama", "Llama 3", "LangChain", "Qdrant", "Docker"],
        "qdrant_collection": COLLECTION_NAME,
        "embedding_model": EMBEDDING_MODEL,
        "vector_size": VECTOR_SIZE,
    }


@app.post("/ingest")
def ingest(payload: IngestRequest):
    docs = [Document(page_content=text, metadata={"source": "manual_input"}) for text in payload.texts if text.strip()]

    if not docs:
        return {"message": "No valid texts provided", "added": 0}

    try:
        vector_store.add_documents(docs)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=(
                "Failed to add documents. Make sure Ollama is running and the embedding model "
                f"'{EMBEDDING_MODEL}' is pulled. Original error: {exc}"
            ),
        )

    return {"message": "Documents added", "added": len(docs)}


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="File name is empty")

    safe_filename = Path(file.filename).name
    file_path = UPLOAD_DIR / safe_filename
    file_path.write_bytes(await file.read())

    text = read_uploaded_file(file_path)

    if not text.strip():
        raise HTTPException(status_code=400, detail="File does not contain readable text")

    docs = split_text_to_documents(text=text, source=safe_filename)

    try:
        vector_store.add_documents(docs)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=(
                "Failed to index uploaded file. Make sure Ollama is running and the embedding model "
                f"'{EMBEDDING_MODEL}' is pulled. Original error: {exc}"
            ),
        )

    return {
        "message": "File uploaded and indexed",
        "filename": safe_filename,
        "chunks": len(docs),
    }


@app.post("/ask")
def ask(payload: AskRequest):
    question = payload.question.strip()

    if not question:
        return {"answer": "Please provide a question.", "sources": []}

    try:
        docs = vector_store.similarity_search(question, k=3)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=(
                "Failed to search documents. Make sure Ollama is running and the embedding model "
                f"'{EMBEDDING_MODEL}' is pulled. Original error: {exc}"
            ),
        )

    if not docs:
        return {
            "answer": "I do not have any indexed documents yet. Add text or upload a file first.",
            "sources": [],
        }

    context = "\n\n".join(doc.page_content for doc in docs)

    prompt = f"""
You are a helpful local AI assistant.
Answer using only the context below.
If the answer is not in the context, say that you do not know.

Context:
{context}

Question:
{question}
"""

    try:
        response = llm.invoke(prompt)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=(
                "Failed to call the LLM. Make sure Ollama is running and the model "
                f"'{LLM_MODEL}' is pulled. Original error: {exc}"
            ),
        )

    return {
        "answer": response.content,
        "sources": [
            {
                "source": doc.metadata.get("source", "unknown"),
                "content": doc.page_content,
            }
            for doc in docs
        ],
    }